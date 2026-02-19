
import os
import sys
from docx import Document

def convert_docx_to_md(docx_path, md_path):
    if not os.path.exists(docx_path):
        print(f"Error: File '{docx_path}' not found.")
        return

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error opening .docx file: {e}")
        return

    md_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Add a blank line for spacing
            md_content.append("")
            continue

        # Basic heading detection based on style name
        style_name = para.style.name.lower()
        if 'heading 1' in style_name:
            md_content.append(f"# {text}")
        elif 'heading 2' in style_name:
            md_content.append(f"## {text}")
        elif 'heading 3' in style_name:
            md_content.append(f"### {text}")
        elif 'list bullet' in style_name:
             md_content.append(f"- {text}")
        elif 'list number' in style_name:
             md_content.append(f"1. {text}") # Simple numbering
        else:
            md_content.append(text)
        
        md_content.append("") # Add newline after paragraph

    # Simple table support (very basic)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            md_content.append("| " + " | ".join(row_cells) + " |")
        md_content.append("")

    try:
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(md_content))
        print(f"Successfully converted '{docx_path}' to '{md_path}'")
    except Exception as e:
        print(f"Error writing to .md file: {e}")

if __name__ == "__main__":
    docx_file = "Assets/website.docx"
    md_file = "Assets/website.md"
    convert_docx_to_md(docx_file, md_file)
